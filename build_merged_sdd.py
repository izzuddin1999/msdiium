from pathlib import Path
import sqlite3

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("SDD-review.docx")
OUTPUT = Path("SDD-HR-19-Policy-and-Guidelines-Merged-Database.docx")
DATABASE = Path("laravel-app/database/database.sqlite")

BUSINESS_TABLES = [
    "users",
    "policy_documents",
    "topic_categories",
    "topic_subtopics",
    "document_histories",
    "document_attachments",
    "document_activity_logs",
    "lookup_values",
    "notifications",
]

MAPPINGS = [
    ("DOCUMENT", "policy_documents", "Core policy, guideline, and circular record; version rows retain parent linkage."),
    ("MAIN_TOPIC", "topic_categories", "Primary controlled document classification."),
    ("SUB_TOPIC", "topic_subtopics", "Secondary classification linked to a main topic."),
    ("DOCUMENT_HISTORY", "document_histories", "Normalized version status, revision, creator, and publication history."),
    ("DOCUMENT_ATTACHMENT", "document_attachments", "Immutable attachment metadata linked to a document and version history."),
    ("DOCUMENT_LOG", "document_activity_logs", "Auditable old/new values, action, user, IP address, and timestamp."),
    ("USER_CAS", "users", "Application identity plus staff ID, CAS username, role, unit, and synchronization timestamp."),
    ("LOV_MAIN", "lookup_values", "Configurable types, statuses, ordering, and active flags."),
    ("NOTIFICATION", "notifications", "Laravel database notification records with payload and read timestamp."),
]


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_table(table, widths):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
                    if row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            if row_index == 0:
                shade(cell, "0B6655")


doc = Document(SOURCE)
section = doc.add_section(WD_SECTION.NEW_PAGE)
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

heading = doc.add_heading("IMPLEMENTED DATABASE MERGE", level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
intro = doc.add_paragraph(
    "This appendix merges the implemented HR-19 Laravel database into the Software Design Description. "
    "It preserves the SDD logical entities while recording the physical table names and columns currently deployed."
)
intro.paragraph_format.space_after = Pt(10)

doc.add_heading("Implementation Mapping", level=2)
mapping_table = doc.add_table(rows=1, cols=3)
for cell, text in zip(mapping_table.rows[0].cells, ["SDD Entity", "Implemented Table", "Implementation Notes"]):
    cell.text = text
for sdd_name, actual_name, notes in MAPPINGS:
    cells = mapping_table.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = sdd_name, actual_name, notes
format_table(mapping_table, [1.7, 2.1, 6.0])

doc.add_paragraph()
doc.add_heading("Physical Schema", level=2)
connection = sqlite3.connect(DATABASE)
connection.row_factory = sqlite3.Row

for table_name in BUSINESS_TABLES:
    exists = connection.execute(
        "SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table_name,)
    ).fetchone()
    if not exists:
        continue

    doc.add_heading(table_name.upper(), level=3)
    columns = connection.execute(f'PRAGMA table_info("{table_name}")').fetchall()
    foreign_keys = connection.execute(f'PRAGMA foreign_key_list("{table_name}")').fetchall()
    fk_map = {row["from"]: f'{row["table"]}.{row["to"]}' for row in foreign_keys}

    table = doc.add_table(rows=1, cols=6)
    headers = ["Column", "Type", "Nullable", "Default", "Key", "Reference"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for column in columns:
        cells = table.add_row().cells
        values = [
            column["name"],
            column["type"] or "-",
            "No" if column["notnull"] else "Yes",
            str(column["dflt_value"]) if column["dflt_value"] is not None else "-",
            "PK" if column["pk"] else ("FK" if column["name"] in fk_map else "-"),
            fk_map.get(column["name"], "-"),
        ]
        for cell, value in zip(cells, values):
            cell.text = value
    format_table(table, [2.0, 1.4, 0.8, 1.7, 0.7, 3.2])
    doc.add_paragraph()

connection.close()

doc.add_heading("Implementation Notes", level=2)
for text in [
    "Laravel plural snake_case names are retained as the physical naming standard.",
    "Document versions remain navigable through policy_documents.parent_document_id and are also represented in document_histories for normalized reporting.",
    "File replacement creates a new document_attachments row so historical attachment evidence is retained.",
    "document_activity_logs stores auditable before-and-after values for governed changes.",
    "CAS integration fields are present, while connection to the institutional CAS service remains an external deployment integration.",
    "Framework operational tables such as cache, jobs, sessions, and migrations are intentionally excluded from this business-schema appendix.",
]:
    paragraph = doc.add_paragraph(style="List Paragraph")
    paragraph.add_run(text)

doc.save(OUTPUT)
print(OUTPUT.resolve())
