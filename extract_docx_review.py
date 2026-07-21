from pathlib import Path
from docx import Document


def extract(path: Path) -> str:
    doc = Document(path)
    lines = [f"# {path.name}", "", "## Paragraphs"]
    for i, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if text:
            lines.append(f"P{i} [{paragraph.style.name}]: {text}")
    lines.extend(["", "## Tables"])
    for table_index, table in enumerate(doc.tables, 1):
        lines.append(f"\n### Table {table_index}")
        for row_index, row in enumerate(table.rows, 1):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            lines.append(f"R{row_index}: " + " | ".join(cells))
    return "\n".join(lines)


for name in ("FRS-review.docx", "SRS-review.docx", "SDD-review.docx"):
    source = Path(name)
    source.with_suffix(".txt").write_text(extract(source), encoding="utf-8")
